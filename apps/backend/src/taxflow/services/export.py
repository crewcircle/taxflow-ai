import io

import markdown as md
from docx import Document
from docx.shared import Cm, Pt
from weasyprint import HTML


def generate_docx(content_md: str, title: str, client_name: str, date: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    header = doc.sections[0].header.paragraphs[0]
    header.text = f"TaxFlow AI | {client_name} | {date}"

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "AI-assisted advice - requires professional review before reliance"

    title_p = doc.add_heading(title, level=1)
    for run in title_p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(14)

    for line in content_md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_pdf(content_md: str, title: str, client_name: str, date: str) -> bytes:
    html_body = md.markdown(content_md)
    html = f"""
    <html>
    <head><style>
      body {{ font-family: Calibri, sans-serif; font-size: 11pt; margin: 2.5cm; }}
      h1 {{ font-size: 14pt; }}
      .header {{ font-size: 9pt; color: #555; }}
      .footer {{ font-size: 8pt; color: #888; }}
    </style></head>
    <body>
      <div class="header">TaxFlow AI | {client_name} | {date}</div>
      <h1>{title}</h1>
      {html_body}
      <div class="footer">AI-assisted advice - requires professional review before reliance</div>
    </body>
    </html>
    """
    return HTML(string=html).write_pdf()
